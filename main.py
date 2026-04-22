#!/usr/bin/env python3
"""
YouTube Channel Summarizer -> Word (.docx)
- Resolves @handles to UC… channel IDs
- Lists latest N video IDs per channel
- Fetches transcripts (youtube-transcript-api) with yt-dlp + Rustypipe-Botguard PO-token provider fallback
- Summarizes with Gemini (Google Generative Language API)
- Exports a .docx report

Run (example):
python main.py --channels channels.txt --out summaries.docx --per-channel 30 --cookies-from-browser firefox --summarize-url-fallback

Optional extras:
  --cookies-from-browser chrome --cookies-profile "Default"
  --cookies cookies.txt
  --rustypipe-bin C:\\path\\to\\rustypipe-botguard.exe   (only if not on PATH)
  --po-token ios.gvs+...    (advanced, only if you actually captured one)
"""

import argparse
import os
import re
import time
import logging
import tempfile
from pathlib import Path
from typing import List, Optional

import requests
from dotenv import load_dotenv
import yt_dlp
from googleapiclient.discovery import build
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound
from docx import Document
from docx.shared import Pt
from datetime import datetime
from googleapiclient.errors import HttpError


# -------------------- Config --------------------

def load_env(require_gemini: bool = True) -> None:
    load_dotenv()
    required_keys = ["YOUTUBE_API_KEY"]
    if require_gemini:
        required_keys.append("GEMINI_API_KEY")

    missing = []
    for key in required_keys:
        if not os.getenv(key):
            missing.append(key)
    if missing:
        raise RuntimeError(f"Missing env vars: {', '.join(missing)}. Create a .env file (see .env.example).")


def build_youtube_client():
    return build("youtube", "v3", developerKey=os.environ["YOUTUBE_API_KEY"])


# -------------------- YouTube helpers --------------------


def resolve_channel_id(youtube, url: str) -> str:
    """Resolve a YouTube channel URL/handle to a UC… channelId."""
    # 1) Direct UC id in /channel/UCxxxx
    m = re.search(r"/channel/(UC[0-9A-Za-z_-]{22})", url)
    if m:
        return m.group(1)

    # 2) Handle /@name  -> use channels().list(forHandle=...)
    h = re.search(r"/@([^/\s]+)", url)
    if h:
        handle = h.group(1)
        handle_param = handle if handle.startswith("@") else f"@{handle}"
        try:
            resp = youtube.channels().list(
                part="id",
                forHandle=handle_param,
            ).execute()
            items = resp.get("items", [])
            if items:
                # in a channel resource, `id` is the channelId
                return items[0]["id"]
        except HttpError as e:
            logging.warning(f"channels.list(forHandle=...) failed for @{handle}: {e}")

        # if forHandle somehow fails, fall back to your old search logic as a last resort
        resp = youtube.search().list(
            q=handle_param,
            type="channel",
            part="snippet",
            maxResults=1,
        ).execute()
        items = resp.get("items", [])
        if items:
            return items[0]["snippet"]["channelId"]
        raise ValueError(f"Could not resolve channel from handle: {handle_param}")

    # 3) Fallback: last path segment as a search term (unchanged)
    guess = url.rstrip("/").split("/")[-1]
    resp = youtube.search().list(q=guess, type="channel", part="snippet", maxResults=1).execute()
    items = resp.get("items", [])
    if items:
        return items[0]["snippet"]["channelId"]
    raise ValueError(f"Could not resolve channel ID from URL: {url}")


def _get_channel_uploads_playlist_id(youtube, channel_id: str) -> Optional[str]:
    """Return channel uploads playlist id (UU...) if available."""
    try:
        resp = youtube.channels().list(part="contentDetails", id=channel_id).execute()
        items = resp.get("items", [])
        if not items:
            return None
        related = items[0].get("contentDetails", {}).get("relatedPlaylists", {})
        return related.get("uploads")
    except Exception:
        return None


def _list_latest_video_ids_from_uploads_playlist(youtube, uploads_playlist_id: str, limit: int) -> List[str]:
    """List latest uploads via playlistItems.list (cheaper quota than search.list)."""
    video_ids: List[str] = []
    page_token = None
    while len(video_ids) < limit:
        resp = youtube.playlistItems().list(
            playlistId=uploads_playlist_id,
            part="snippet",
            maxResults=min(50, limit - len(video_ids)),
            pageToken=page_token,
        ).execute()
        items = resp.get("items", [])
        for item in items:
            vid = (item.get("snippet", {}).get("resourceId", {}) or {}).get("videoId")
            if vid:
                video_ids.append(vid)
        page_token = resp.get("nextPageToken")
        if not page_token:
            break
    return video_ids[:limit]


def _list_latest_video_ids_search(youtube, channel_id: str, limit: int = 30) -> List[str]:
    video_ids: List[str] = []
    page_token = None
    while len(video_ids) < limit:
        req = youtube.search().list(
            channelId=channel_id,
            order="date",
            type="video",
            part="id",
            maxResults=min(50, limit - len(video_ids)),
            pageToken=page_token,
        )
        resp = req.execute()
        for item in resp.get("items", []):
            vid = item["id"].get("videoId")
            if vid:
                video_ids.append(vid)
        page_token = resp.get("nextPageToken")
        if not page_token:
            break
    return video_ids[:limit]


def list_latest_video_ids(youtube, channel_id: str, limit: int = 30) -> List[str]:
    uploads_playlist_id = _get_channel_uploads_playlist_id(youtube, channel_id)
    if uploads_playlist_id:
        return _list_latest_video_ids_from_uploads_playlist(youtube, uploads_playlist_id, limit)
    return _list_latest_video_ids_search(youtube, channel_id, limit)


def get_video_snippets_batch(youtube, video_ids: List[str]) -> dict:
    """Fetch snippets for many video ids using batched videos.list calls."""
    out = {}
    if not video_ids:
        return out
    # De-duplicate while preserving order
    unique_video_ids = list(dict.fromkeys(video_ids))
    for i in range(0, len(unique_video_ids), 50):
        chunk = unique_video_ids[i:i + 50]
        try:
            resp = youtube.videos().list(part="snippet", id=",".join(chunk)).execute()
            for item in resp.get("items", []):
                vid = item.get("id")
                if vid:
                    out[vid] = item.get("snippet", {}) or {}
        except Exception:
            pass
    return out


def get_video_snippet(youtube, video_id: str) -> dict:
    """Fetch snippet for a single id."""
    snippets = get_video_snippets_batch(youtube, [video_id])
    return snippets.get(video_id, {})


def transcript_file_exists(video_id: str, title: str, published_at: Optional[str]) -> bool:
    """Check whether transcript file already exists for this video/title/date tuple."""
    transcripts_dir = Path("transcripts")
    short_title = create_short_title(title)
    date_str = format_date_yymmdd(published_at) or "unknown"
    filename = f"{short_title}_{video_id}_{date_str}.txt"
    return (transcripts_dir / filename).exists()


def get_video_id_from_transcript_filename(filename: str) -> Optional[str]:
    """Extract the 11-char YouTube video id from transcript filename."""
    m = re.search(r"_([0-9A-Za-z_-]{11})_(\d{6}|unknown)\.txt$", filename)
    if m:
        return m.group(1)
    return None


def get_existing_transcript_video_ids() -> set:
    """Collect video ids from existing transcript files."""
    transcripts_dir = Path("transcripts")
    if not transcripts_dir.exists():
        return set()
    video_ids = set()
    try:
        for path in transcripts_dir.glob("*.txt"):
            vid = get_video_id_from_transcript_filename(path.name)
            if vid:
                video_ids.add(vid)
    except Exception:
        pass
    return video_ids


def format_published_date(iso_str: Optional[str], fmt: str) -> Optional[str]:
    """Format ISO 8601 timestamp (e.g., 2024-05-21T12:34:56Z) to the given strftime format."""
    if not iso_str:
        return None
    try:
        # Handle trailing Z / UTC
        if iso_str.endswith("Z"):
            iso_str = iso_str.replace("Z", "+00:00")
        dt = datetime.fromisoformat(iso_str)
        return dt.strftime(fmt)
    except Exception:
        return iso_str  # fall back to raw


def create_short_title(title: str, max_words: int = 6) -> str:
    """Create a short title from the video title (5-6 words, separated by underscores)."""
    if not title:
        return "untitled"
    # Remove special characters and split into words
    words = re.sub(r'[^\w\s]', ' ', title).split()
    # Take first max_words words
    short_words = words[:max_words]
    # Join with underscores and make filesystem-safe
    short_title = "_".join(short_words).lower()
    # Remove any remaining invalid characters
    short_title = re.sub(r'[^\w_-]', '', short_title)
    return short_title if short_title else "untitled"


def format_date_yymmdd(iso_str: Optional[str]) -> Optional[str]:
    """Format ISO 8601 timestamp to YYMMDD format."""
    if not iso_str:
        return None
    try:
        # Handle trailing Z / UTC
        if iso_str.endswith("Z"):
            iso_str = iso_str.replace("Z", "+00:00")
        dt = datetime.fromisoformat(iso_str)
        return dt.strftime("%y%m%d")
    except Exception:
        return None


def save_transcript(transcript: str, video_id: str, title: str, published_at: Optional[str]) -> None:
    """Save transcript to transcripts folder with format: [short_title]_[youtube_id]_[YYMMDD].txt"""
    transcripts_dir = Path("transcripts")
    transcripts_dir.mkdir(exist_ok=True)
    
    short_title = create_short_title(title)
    date_str = format_date_yymmdd(published_at) or "unknown"
    
    filename = f"{short_title}_{video_id}_{date_str}.txt"
    filepath = transcripts_dir / filename
    
    try:
        filepath.write_text(transcript, encoding="utf-8")
        logging.info(f"💾 Saved transcript: {filename}")
    except Exception as e:
        logging.warning(f"⚠️  Failed to save transcript {filename}: {e}")


# -------------------- Transcript fetch (yt-transcript + yt-dlp fallback) --------------------

def _strip_vtt(vtt_text: str) -> str:
    lines = []
    for line in vtt_text.splitlines():
        if (not line) or line.startswith("WEBVTT") or re.match(r"^\d+$", line) or re.search(r"-->\s", line):
            continue
        lines.append(line.strip())
    return " ".join(lines)


def fetch_transcript_with_ytdlp(video_id: str, lang: str = "en") -> Optional[str]:
    """Try yt-dlp with Rustypipe-Botguard provider to fetch auto-subs as VTT and return plain text."""
    url = f"https://youtu.be/{video_id}"
    tmpdir = Path(tempfile.mkdtemp(prefix="ytdlp_") )
    outtmpl = str(tmpdir / "%(id)s.%(ext)s")

    # Base opts
    ydl_opts = {
        "skip_download": True,
        "writesubtitles": True,
        "writeautomaticsub": True,
        "subtitleslangs": [lang, "en", "en-US", "en-GB"],
        "subtitlesformat": "vtt",
        "outtmpl": outtmpl,
        "quiet": True,
        "nocheckcertificate": True,
    }

    ydl_opts["verbose"] = True

    # Extractor args (PO-token provider + client choice)
    ea = ydl_opts.setdefault("extractor_args", {}).setdefault("youtube", {})
    ea["pot_providers"] = ["rustypipe-botguard"]
    ea["rustypipe_bg_pot_cache"] = ["1"]  # use provider cache

    cookies_path = getattr(fetch_transcript_with_ytdlp, "_cookies_path", None)
    cookies_browser = getattr(fetch_transcript_with_ytdlp, "_cookies_browser", None)
    cookies_profile = getattr(fetch_transcript_with_ytdlp, "_cookies_profile", None)
    po_token = getattr(fetch_transcript_with_ytdlp, "_po_token", None)
    rp_bin = getattr(fetch_transcript_with_ytdlp, "_rustypipe_bin", None) or os.getenv("RUSTYPIPE_BG_BIN")

    if rp_bin:
        ea["rustypipe_bg_bin"] = [rp_bin]

    # Client selection
    if po_token:
        ea["player_client"] = ["ios", "web_safari", "web"]
        ea["po_token"] = [po_token]
    else:
        if cookies_browser or cookies_path:
            ea["player_client"] = ["web_safari", "web"]
        else:
            ea["player_client"] = ["android", "web_safari", "web"]

    # Cookies hookup
    if cookies_path:
        ydl_opts["cookiefile"] = cookies_path
    if cookies_browser:
        ydl_opts["cookiesfrombrowser"] = (cookies_browser,) if not cookies_profile else (cookies_browser, cookies_profile)

    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.extract_info(url, download=True)
        # Look for any VTT that yt-dlp produced for this video (e.g., id.vtt or id.en.vtt)
        vtt_candidates = list(tmpdir.glob(f"{video_id}*.vtt"))
        if vtt_candidates:
            # Prefer the largest (usually the full auto-cc track)
            vtt_path = max(vtt_candidates, key=lambda p: p.stat().st_size)
            vtt_text = vtt_path.read_text(encoding="utf-8", errors="ignore")
            text = _strip_vtt(vtt_text)
            if text.strip():
                return text

    except Exception as e:
        logging.warning(f"yt-dlp fallback failed for {video_id}: {e}")
        return None
    finally:
        # cleanup
        try:
            for f in tmpdir.glob("*"):
                f.unlink()
            tmpdir.rmdir()
        except Exception:
            pass
    return None


def fetch_transcript_text(video_id: str) -> Optional[str]:
    """First try youtube-transcript-api; on failure, try yt-dlp fallback."""
    try:
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
        transcript = None
        for lang in ("en", "en-US", "en-GB"):
            try:
                transcript = transcript_list.find_transcript([lang])
                if transcript:
                    break
            except Exception:
                pass
        if not transcript:
            for t in transcript_list:
                if getattr(t, 'is_generated', False) or str(getattr(t, 'language_code', '')).startswith("en"):
                    transcript = t
                    break
        if transcript:
            chunks = transcript.fetch()
            return " ".join(c.get("text", "") for c in chunks if c.get("text"))
    except (TranscriptsDisabled, NoTranscriptFound):
        pass
    except Exception:
        pass
    return fetch_transcript_with_ytdlp(video_id, os.getenv("SUB_LANG", "en"))


# -------------------- Gemini --------------------

def _gemini_request(payload: dict) -> str:
    api_key = os.environ["GEMINI_API_KEY"]
    endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{os.getenv('GEMINI_MODEL', 'gemini-1.5-flash')}:generateContent"
    headers = {"Content-Type": "application/json", "x-goog-api-key": api_key}
    r = requests.post(endpoint, headers=headers, json=payload, timeout=90)
    r.raise_for_status()
    data = r.json()
    try:
        return data["candidates"][0]["content"]["parts"][0]["text"].strip()
    except Exception:
        return "Summary unavailable (unexpected API response)."


def gemini_summarize(transcript: str) -> str:
    prompt = (
        "You are a helpful assistant. Summarize this YouTube video transcript in:\n"
        "- a 3–5 sentence overview,\n"
        "- 5 concise bullet points of key insights,\n"
        "- and 3 suggested action items if applicable.\n\n"
        "Keep it under 1800 characters total. Here is the transcript:\n\n"
    )
    return _gemini_request({"contents": [{"parts": [{"text": prompt + transcript[:25000]}]}]})


def gemini_summarize_url(video_url: str) -> str:
    prompt = (
        "Given this YouTube link, provide a concise summary based on any metadata or content you can access. "
        "If you cannot access the video, say so explicitly and provide a one-sentence guess based on the title only.\n"
        f"URL: {video_url}"
    )
    return _gemini_request({"contents": [{"parts": [{"text": prompt}]}]})


def gemini_summarize_metadata(meta: dict) -> str:
    """Fallback to summarize title+description when no transcript/URL content is available."""
    blob = (
        f"Title: {meta.get('title','')}\n"
        f"Channel: {meta.get('channelTitle','')}\n"
        f"Published: {meta.get('publishedAt','')}\n\n"
        f"Description:\n{meta.get('description','')}\n"
    )
    prompt = "Summarize the following video metadata as best as possible:\n\n" + blob
    return _gemini_request({"contents": [{"parts": [{"text": prompt[:25000]}]}]})


# -------------------- Document helpers --------------------

def add_channel_heading(doc: Document, channel_url: str):
    h = doc.add_heading(level=1)
    run = h.add_run(f"Channel: {channel_url}")
    run.font.size = Pt(16)


def add_video_section(doc: Document, video_id: str, summary: str, published: Optional[str] = None):
    title = f"Video: https://youtu.be/{video_id}"
    if published:
        title += f" — Published: {published}"
    doc.add_heading(title, level=2)
    doc.add_paragraph(summary)


# -------------------- Main --------------------

def main():
    parser = argparse.ArgumentParser(description="Summarize latest YouTube videos with Gemini and export to DOCX.")
    parser.add_argument("--channels", type=str, required=True, help="Path to a text file with one YouTube channel URL per line.")
    parser.add_argument("--out", type=str, default="YouTube_Summaries.docx", help="Output .docx file path.")
    parser.add_argument("--transcripts-only", action="store_true",
                        help="Fetch and save transcripts only; skip Gemini summarization and DOCX output.")
    parser.add_argument("--skip-existing", action=argparse.BooleanOptionalAction, default=True,
                        help="Skip videos that already have transcript files (default: true).")
    parser.add_argument("--per-channel", type=int, default=30, help="How many latest videos to fetch per channel (default 30).")
    parser.add_argument("--sleep", type=float, default=0.5, help="Sleep (seconds) between API calls.")
    parser.add_argument("--summarize-url-fallback", action="store_true", help="If no transcript, ask Gemini to summarize the video URL directly.")
    parser.add_argument("--sub-lang", type=str, default="en", help="Preferred subtitle language for yt-dlp fallback (default: en).")
    parser.add_argument("--cookies", type=str, default=None, help="Path to a cookies.txt file for YouTube (netscape format).")
    parser.add_argument("--cookies-from-browser", type=str, default=None, help="Load cookies from a local browser (e.g., chrome, firefox, edge).")
    parser.add_argument("--cookies-profile", type=str, default=None, help="Profile name for --cookies-from-browser (e.g. 'Default').")
    parser.add_argument("--rustypipe-bin", type=str, default=os.getenv("RUSTYPIPE_BG_BIN"), help="Path to rustypipe-botguard (if not on PATH).")
    parser.add_argument("--po-token", type=str, default=os.getenv("YT_PO_TOKEN"), help="GVS PO token (advanced; enables iOS client).")

    # NEW: include published date in headings
    parser.add_argument("--show-date", action="store_true",
                        help="Include each video's published date in the DOCX heading.")
    parser.add_argument("--date-format", type=str, default="%Y-%m-%d",
                        help="strftime format for dates when --show-date is used (default: %%Y-%%m-%%d).")

    args = parser.parse_args()
    load_env(require_gemini=not args.transcripts_only)
    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
    log = logging.getLogger(__name__)

    # Wire values for yt-dlp helper
    fetch_transcript_with_ytdlp._cookies_path = args.cookies
    fetch_transcript_with_ytdlp._cookies_browser = args.cookies_from_browser
    fetch_transcript_with_ytdlp._cookies_profile = args.cookies_profile
    fetch_transcript_with_ytdlp._rustypipe_bin = args.rustypipe_bin
    fetch_transcript_with_ytdlp._po_token = args.po_token

    youtube = build_youtube_client()
    doc = None
    if not args.transcripts_only:
        doc = Document()
        doc.add_heading("YouTube Channel Summaries", level=0)

    with open(args.channels, "r", encoding="utf-8") as f:
        channel_urls = [line.strip() for line in f if line.strip()]

    for ch_url in channel_urls:
        log.info(f"�� Processing channel: {ch_url}")
        if doc is not None:
            add_channel_heading(doc, ch_url)
        try:
            ch_id = resolve_channel_id(youtube, ch_url)
            log.info(f"✅ Resolved channel ID: {ch_id}")
        except Exception as e:
            log.error(f"❌ Could not resolve channel: {e}")
            if doc is not None:
                doc.add_paragraph(f"Could not resolve channel: {e}")
            continue

        video_ids = list_latest_video_ids(youtube, ch_id, args.per_channel)
        log.info(f"�� Found {len(video_ids)} videos for {ch_url}")
        snippets_by_id = get_video_snippets_batch(youtube, video_ids)
        existing_video_ids = get_existing_transcript_video_ids() if args.skip_existing else set()

        for vid in video_ids:
            # Always fetch snippet to get title and published date (needed for transcript saving)
            snippet_for_meta = snippets_by_id.get(vid, {})
            video_title = snippet_for_meta.get("title", "")
            published_at = snippet_for_meta.get("publishedAt")

            if args.skip_existing and vid in existing_video_ids:
                log.info(f"ℹ️  Skipping existing transcript for https://youtu.be/{vid}")
                time.sleep(args.sleep)
                continue
            
            # Format published date for display if requested
            published_str = None
            if args.show_date:
                published_str = format_published_date(published_at, args.date_format)

            log.info(f"�� Fetching transcript for https://youtu.be/{vid}")
            transcript = fetch_transcript_text(vid)

            if not transcript:
                msg = "No transcript available."
                if args.transcripts_only:
                    log.info(f"ℹ️  {msg} for https://youtu.be/{vid}")
                    time.sleep(args.sleep)
                    continue
                summary_done = False
                if args.summarize_url_fallback:
                    log.info("No transcript. Trying URL summarization fallback…")
                    url = f"https://youtu.be/{vid}"
                    try:
                        summary = gemini_summarize_url(url)
                        if summary and len(summary.strip()) > 40:
                            add_video_section(doc, vid, summary, published=published_str)
                            summary_done = True
                    except Exception as e:
                        msg += f" URL fallback error: {e}"

                if not summary_done:
                    # Metadata fallback: title+description
                    try:
                        meta = {
                            "title": snippet_for_meta.get("title", ""),
                            "description": snippet_for_meta.get("description", ""),
                            "publishedAt": snippet_for_meta.get("publishedAt", ""),
                            "channelTitle": snippet_for_meta.get("channelTitle", ""),
                        }
                        summary = gemini_summarize_metadata(meta)
                        add_video_section(doc, vid, summary, published=published_str)
                        summary_done = True
                    except Exception:
                        pass

                if not summary_done:
                    add_video_section(doc, vid, msg, published=published_str)
                time.sleep(args.sleep)
                continue

            # Save transcript to file
            save_transcript(transcript, vid, video_title, published_at)
            if args.transcripts_only:
                time.sleep(args.sleep)
                continue

            try:
                log.info("�� Summarizing with Gemini…")
                summary = gemini_summarize(transcript)
                log.info("✅ Summary complete.")
            except Exception as e:
                log.error(f"❌ Gemini error: {e}")
                summary = f"Error summarizing: {e}"
            add_video_section(doc, vid, summary, published=published_str)
            time.sleep(args.sleep)

    if doc is not None:
        doc.save(args.out)
        print(f"✅ Done! File saved as: {args.out}")
    else:
        print("✅ Done! Transcripts saved in: transcripts/")


if __name__ == "__main__":
    main()
